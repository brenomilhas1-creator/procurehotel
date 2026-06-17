"""
Serviço de OCR / extração de tabelas.

Usa Docling como motor principal (qualidade superior para tabelas em PDF/Word)
e Tesseract como fallback. Para Excel/CSV usa pandas.
"""

from __future__ import annotations

import csv
import io
import os
from pathlib import Path
from typing import Any

from loguru import logger

from app.core.config import settings


def _ensure_dir(p: str) -> str:
    Path(p).mkdir(parents=True, exist_ok=True)
    return p


async def extract_rows(file_path: str, mime_type: str | None) -> list[dict[str, Any]]:
    """Deteta o tipo de ficheiro e devolve lista de {row_index, raw_text, cells}."""
    p = Path(file_path)
    suffix = p.suffix.lower()
    name = p.name.lower()

    if suffix in {".csv"} or "csv" in (mime_type or ""):
        return _from_csv(p)
    if suffix in {".xlsx", ".xlsm", ".xls"} or "spreadsheet" in (mime_type or ""):
        return _from_excel(p)
    if suffix in {".pdf"} or "pdf" in (mime_type or ""):
        return await _from_pdf(p)
    if suffix in {".docx", ".doc"} or "word" in (mime_type or ""):
        return _from_docx(p)
    if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}:
        return await _from_image(p)

    raise ValueError(f"Tipo de ficheiro não suportado: {suffix or mime_type}")


# ---------------------------------------------------------------------------
# CSV / Excel / Word
# ---------------------------------------------------------------------------


def _from_csv(p: Path) -> list[dict]:
    rows: list[dict] = []
    with p.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f)
        for i, cells in enumerate(reader):
            if not cells or all((c or "").strip() == "" for c in cells):
                continue
            rows.append({
                "row_index": i,
                "raw_text": " | ".join((c or "") for c in cells),
                "cells": [c.strip() if c else None for c in cells],
            })
    return rows


def _from_excel(p: Path) -> list[dict]:
    import openpyxl
    wb = openpyxl.load_workbook(p, data_only=True, read_only=True)
    rows: list[dict] = []
    for ws in wb.worksheets:
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            cells = [(str(c).strip() if c is not None else None) for c in row]
            if not any(c for c in cells):
                continue
            rows.append({
                "row_index": i,
                "raw_text": " | ".join((c or "") for c in cells),
                "cells": cells,
            })
    return rows


def _from_docx(p: Path) -> list[dict]:
    from docx import Document
    doc = Document(str(p))
    rows: list[dict] = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells = [(c.text.strip() or None) for c in row.cells]
            if not any(cells):
                continue
            rows.append({
                "row_index": f"{t_idx}.{r_idx}",
                "raw_text": " | ".join((c or "") for c in cells),
                "cells": cells,
            })
    if not rows:
        # fallback: parágrafos
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            rows.append({
                "row_index": i,
                "raw_text": text,
                "cells": [text],
            })
    return rows


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


async def _from_pdf(p: Path) -> list[dict]:
    if settings.ocr_engine == "docling":
        try:
            return _from_docling(str(p))
        except Exception as e:
            logger.warning(f"Docling falhou, a usar pypdf: {e}")

    # Fallback com pypdf + pytesseract em imagens embebidas
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(p))
        rows: list[dict] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            for line in text.splitlines():
                line = line.strip()
                if not line:
                    continue
                rows.append({
                    "row_index": f"{i}.{len(rows)}",
                    "raw_text": line,
                    "cells": [seg.strip() for seg in line.split("  ") if seg.strip()],
                })
        return rows
    except Exception as e:
        logger.error(f"Falha no PDF fallback: {e}")
        return []


def _from_docling(path: str) -> list[dict]:
    """Usa Docling para extrair tabelas de um PDF/DOCX."""
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        raise RuntimeError("docling não instalado; a usar fallback")
    converter = DocumentConverter()
    result = converter.convert(path)
    rows: list[dict] = []
    idx = 0
    for table in (result.document.tables or []):
        df = table.export_to_dataframe()
        for _, r in df.iterrows():
            cells = [str(v).strip() if v is not None else None for v in r.tolist()]
            rows.append({
                "row_index": idx,
                "raw_text": " | ".join((c or "") for c in cells),
                "cells": cells,
            })
            idx += 1
    if not rows:
        # sem tabelas detetadas: extrair texto
        for i, line in enumerate((result.document.export_to_text() or "").splitlines()):
            line = line.strip()
            if not line:
                continue
            rows.append({
                "row_index": i,
                "raw_text": line,
                "cells": [line],
            })
    return rows


# ---------------------------------------------------------------------------
# Imagem
# ---------------------------------------------------------------------------


async def _from_image(p: Path) -> list[dict]:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(p)
        text = pytesseract.image_to_string(img, lang=settings.tesseract_lang)
        rows: list[dict] = []
        for i, line in enumerate(text.splitlines()):
            line = line.strip()
            if not line:
                continue
            rows.append({
                "row_index": i,
                "raw_text": line,
                "cells": [seg.strip() for seg in line.split("  ") if seg.strip()],
            })
        return rows
    except Exception as e:
        logger.error(f"Falha Tesseract: {e}")
        return []


def save_upload(upload_dir: str, filename: str, content: bytes) -> str:
    """Guarda conteúdo em disco e devolve path absoluto."""
    _ensure_dir(upload_dir)
    safe = os.path.basename(filename).replace("/", "_")
    p = Path(upload_dir) / f"{os.urandom(4).hex()}_{safe}"
    p.write_bytes(content)
    return str(p)
